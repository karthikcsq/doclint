"""DOCX (Word) parser."""

from pathlib import Path
from typing import ClassVar

import docx
from docx.opc.exceptions import PackageNotFoundError

from ..core.document import DocumentMetadata
from ..core.exceptions import ParsingError
from .base import BaseParser


class DOCXParser(BaseParser):
    """Parser for Microsoft Word (.docx) documents.

    Extracts text from paragraphs and tables.
    Extracts metadata from document core properties.
    """

    file_type: ClassVar[str] = "docx"
    supported_extensions: ClassVar[list[str]] = [".docx"]

    def parse(self, path: Path) -> str:
        """Extract text from DOCX file.

        Extracts text from both paragraphs and tables.

        Args:
            path: Path to DOCX file

        Returns:
            Extracted text content

        Raises:
            ParsingError: If file is not a valid DOCX or parsing fails
        """
        try:
            doc = docx.Document(str(path))

            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)

            full_text = "\n\n".join(paragraphs)
            return self.clean_text(full_text)

        except PackageNotFoundError as e:
            raise ParsingError(
                f"Invalid DOCX file {path}: File may be corrupted or not a valid DOCX",
                path=str(path),
                original_error=e,
            )
        except Exception as e:
            raise ParsingError(
                f"Failed to parse DOCX {path}: {e}",
                path=str(path),
                original_error=e,
            )

    def extract_metadata(self, path: Path) -> DocumentMetadata:
        """Extract metadata from DOCX file.

        Tries to extract:
        - Title from document title property
        - Author from document author property
        - Created/modified dates from core properties
        - Version from revision property
        - Keywords as tags

        Args:
            path: Path to DOCX file

        Returns:
            DocumentMetadata
        """
        metadata = DocumentMetadata()

        try:
            doc = docx.Document(str(path))
            core_props = doc.core_properties

            # Extract standard properties
            if core_props.author:
                metadata.author = core_props.author

            if core_props.title:
                metadata.title = core_props.title

            if core_props.created:
                metadata.created = core_props.created

            if core_props.modified:
                metadata.modified = core_props.modified

            if core_props.revision:
                metadata.version = str(core_props.revision)

            # Extract keywords as tags
            if core_props.keywords:
                metadata.tags = [k.strip() for k in core_props.keywords.split(",")]

            # Additional custom properties
            if core_props.subject:
                metadata.custom["subject"] = core_props.subject

            if core_props.category:
                metadata.custom["category"] = core_props.category

            # If no title from metadata, use filename
            if not metadata.title:
                metadata.title = path.stem

        except Exception:
            # Best-effort
            pass

        return metadata

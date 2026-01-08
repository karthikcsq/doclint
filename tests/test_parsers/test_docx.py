"""Tests for DOCX parser."""

from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from doclint.core.document import DocumentMetadata
from doclint.core.exceptions import ParsingError
from doclint.parsers.docx import DOCXParser


class TestDOCXParser:
    """Test suite for DOCXParser."""

    def test_can_parse_docx_files(self) -> None:
        """Test parser recognizes DOCX extension."""
        parser = DOCXParser()
        assert parser.can_parse(Path("test.docx")) is True

    def test_cannot_parse_other_files(self) -> None:
        """Test parser rejects non-DOCX files."""
        parser = DOCXParser()
        assert parser.can_parse(Path("test.txt")) is False
        assert parser.can_parse(Path("test.pdf")) is False

    def test_parse_simple_docx(self, tmp_path: Path) -> None:
        """Test parsing simple DOCX file."""
        test_file = tmp_path / "test.docx"

        # Create a simple DOCX
        doc = Document()
        doc.add_paragraph("Hello, World!")
        doc.add_paragraph("This is a test document.")
        doc.save(str(test_file))

        parser = DOCXParser()
        result = parser.parse(test_file)

        assert "Hello, World!" in result
        assert "test document" in result

    def test_parse_with_table(self, tmp_path: Path) -> None:
        """Test parsing DOCX with tables."""
        test_file = tmp_path / "test.docx"

        # Create DOCX with table
        doc = Document()
        doc.add_paragraph("Before table")

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Cell 1"
        table.cell(0, 1).text = "Cell 2"
        table.cell(1, 0).text = "Cell 3"
        table.cell(1, 1).text = "Cell 4"

        doc.add_paragraph("After table")
        doc.save(str(test_file))

        parser = DOCXParser()
        result = parser.parse(test_file)

        assert "Before table" in result
        assert "Cell 1" in result
        assert "Cell 2" in result
        assert "Cell 3" in result
        assert "Cell 4" in result
        assert "After table" in result

    def test_parse_skips_empty_paragraphs(self, tmp_path: Path) -> None:
        """Test that empty paragraphs are skipped."""
        test_file = tmp_path / "test.docx"

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("")  # Empty paragraph
        doc.add_paragraph("Second paragraph")
        doc.save(str(test_file))

        parser = DOCXParser()
        result = parser.parse(test_file)

        # Should have content from both paragraphs
        assert "First paragraph" in result
        assert "Second paragraph" in result
        # Excessive whitespace should be cleaned
        assert "\n\n\n\n" not in result

    def test_parse_nonexistent_file(self) -> None:
        """Test parsing non-existent file raises error."""
        parser = DOCXParser()

        with pytest.raises(ParsingError):
            parser.parse(Path("/nonexistent/file.docx"))

    def test_parse_invalid_docx(self, tmp_path: Path) -> None:
        """Test parsing invalid DOCX file raises error."""
        test_file = tmp_path / "invalid.docx"
        test_file.write_text("Not a DOCX file", encoding="utf-8")

        parser = DOCXParser()

        with pytest.raises(ParsingError) as exc_info:
            parser.parse(test_file)

        assert "Invalid DOCX" in str(exc_info.value) or "Failed to parse" in str(exc_info.value)

    def test_extract_metadata(self, tmp_path: Path) -> None:
        """Test metadata extraction from DOCX."""
        test_file = tmp_path / "test.docx"

        doc = Document()
        doc.add_paragraph("Content")

        # Set core properties
        core_props = doc.core_properties
        core_props.title = "Test Document"
        core_props.author = "Jane Doe"
        core_props.subject = "Testing"
        core_props.keywords = "test, docx, parser"

        doc.save(str(test_file))

        parser = DOCXParser()
        metadata = parser.extract_metadata(test_file)

        assert metadata.title == "Test Document"
        assert metadata.author == "Jane Doe"
        assert metadata.custom.get("subject") == "Testing"
        assert "test" in metadata.tags
        assert "docx" in metadata.tags
        assert "parser" in metadata.tags

    def test_extract_metadata_with_dates(self, tmp_path: Path) -> None:
        """Test metadata extraction with dates."""
        test_file = tmp_path / "test.docx"

        doc = Document()
        doc.add_paragraph("Content")

        # Set dates
        core_props = doc.core_properties
        core_props.created = datetime(2024, 1, 1, 10, 0, 0)
        core_props.modified = datetime(2024, 2, 15, 14, 30, 0)

        doc.save(str(test_file))

        parser = DOCXParser()
        metadata = parser.extract_metadata(test_file)

        assert metadata.created is not None
        assert metadata.created.year == 2024
        assert metadata.created.month == 1

        assert metadata.modified is not None
        assert metadata.modified.year == 2024
        assert metadata.modified.month == 2

    def test_extract_metadata_fallback_to_filename(self, tmp_path: Path) -> None:
        """Test metadata falls back to filename if no title."""
        test_file = tmp_path / "my-document.docx"

        doc = Document()
        doc.add_paragraph("Content")
        doc.save(str(test_file))

        parser = DOCXParser()
        metadata = parser.extract_metadata(test_file)

        assert metadata.title == "my-document"

    def test_extract_metadata_best_effort(self) -> None:
        """Test metadata extraction doesn't fail on error."""
        parser = DOCXParser()

        # Non-existent file should return empty metadata (not raise)
        metadata = parser.extract_metadata(Path("/nonexistent/file.docx"))

        assert isinstance(metadata, DocumentMetadata)

    def test_file_type_is_docx(self) -> None:
        """Test parser has correct file_type."""
        parser = DOCXParser()
        assert parser.file_type == "docx"

    def test_supported_extensions(self) -> None:
        """Test parser lists supported extensions."""
        parser = DOCXParser()
        assert ".docx" in parser.supported_extensions
